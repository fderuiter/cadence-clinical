import hashlib
import hmac
import json
import os
import re
import time

import pytest
from fastapi.testclient import TestClient

from apps.designer.db import MOCK_DESIGNER_AUDIT_LOGS
from apps.designer.main import app as designer_app
from apps.designer.rendering import (
    TemplateRenderingError,
    build_docx_template,
    get_safe_filename,
    sanitize_filename,
)


def get_custom_auth_headers(change_reason="system_operation", roles="admin"):
    timestamp = str(time.time())
    user_id = "123"
    secret = "internal-gateway-secret-12345"  # pragma: allowlist secret
    payload = {
        "change_reason": change_reason,
        "roles": roles,
        "timestamp": timestamp,
        "user_id": user_id,
    }
    serialized = json.dumps(payload, sort_keys=True, separators=(",", ":"))
    signature = hmac.new(
        secret.encode(), serialized.encode(), hashlib.sha256
    ).hexdigest()
    return {
        "X-User-Id": user_id,
        "X-User-Roles": roles,
        "X-Gateway-Timestamp": timestamp,
        "X-Gateway-Signature": signature,
        "X-Signature-Version": "2",
        "X-Change-Reason": change_reason,
    }


@pytest.fixture
def client():
    return TestClient(designer_app, raise_server_exceptions=False)


@pytest.fixture(autouse=True)
def clear_audit_logs():
    MOCK_DESIGNER_AUDIT_LOGS.clear()
    yield
    MOCK_DESIGNER_AUDIT_LOGS.clear()


@pytest.fixture(scope="module", autouse=True)
def setup_default_template():
    """
    Ensure a baseline protocol template exists for standard tests.
    """
    build_docx_template()


def test_sanitize_filename():
    """
    Verify that sanitize_filename filters out dangerous or unsupported characters.
    """
    assert sanitize_filename("study_123") == "study_123"
    assert sanitize_filename("study/123/../../etc") == "study_123_.._.._etc"
    assert sanitize_filename("My Protocol #1!") == "My_Protocol_1"
    assert sanitize_filename("space test   and symbols @#$") == "space_test_and_symbols"


def test_get_safe_filename():
    """
    Verify safe and deterministic suggested filename derivation.
    """
    filename = get_safe_filename("study/123", 4, "pdf")
    assert filename == "protocol_study_123_v4.pdf"

    filename_docx = get_safe_filename(" oncology-trial ", 1, "docx")
    assert filename_docx == "protocol_oncology-trial_v1.docx"


def test_build_docx_template():
    """
    Verify that build_docx_template successfully generates a file on disk.

    Requirements: PRD-SYS-001
    """
    path = build_docx_template()
    assert os.path.exists(path)
    assert path.endswith(".docx")


def test_template_immutability(tmp_path, monkeypatch):
    """
    Assert that the template on disk remains byte-for-byte unchanged after render_protocol_to_docx().

    Requirements: PRD-SYS-001
    """
    import usdm_model

    from apps.designer import rendering
    from apps.designer.content_assembly import assemble_rendered_protocol_document
    from apps.designer.db import get_study_projection
    from apps.designer.mapper import map_study_to_usdm, to_uuid

    # Isolate TEMPLATES_DIR to a temporary path
    temp_dir = str(tmp_path)
    monkeypatch.setattr(rendering, "TEMPLATES_DIR", temp_dir)

    # Build template in isolated TEMPLATES_DIR
    template_path = rendering.build_docx_template()
    assert os.path.exists(template_path)

    # Capture initial bytes
    with open(template_path, "rb") as f:
        initial_bytes = f.read()

    # Prepare document data for rendering
    study_data = get_study_projection("study_1")
    usdm_dict = map_study_to_usdm(study_data)
    usdm_dict["id"] = to_uuid(usdm_dict["id"], "study")
    study_obj = usdm_model.Study.model_validate(usdm_dict)
    doc_view = assemble_rendered_protocol_document(
        study=study_obj,
        creator="test_user",
        change_reason="Test template rendering",
        version_index=1,
    )

    # Run render
    result = rendering.render_protocol_to_docx(doc_view, "combined")
    assert result.content is not None

    # Assert template file is byte-for-byte unchanged
    with open(template_path, "rb") as f:
        after_bytes = f.read()
    assert after_bytes == initial_bytes


def test_load_template_missing(tmp_path, monkeypatch):
    """
    Assert that trying to load with no template present raises TemplateRenderingError.

    Requirements: PRD-SYS-001
    """
    from apps.designer import rendering

    # Isolate TEMPLATES_DIR to an empty directory
    temp_dir = str(tmp_path)
    monkeypatch.setattr(rendering, "TEMPLATES_DIR", temp_dir)

    with pytest.raises(TemplateRenderingError) as exc_info:
        rendering.load_docx_template()
    assert "Template file is missing" in str(exc_info.value)


def test_load_template_invalid(tmp_path, monkeypatch):
    """
    Assert that an invalid or corrupt template file (e.g. non-zip) raises TemplateRenderingError.

    Requirements: PRD-SYS-001
    """
    from apps.designer import rendering

    # Isolate TEMPLATES_DIR to a temporary path
    temp_dir = str(tmp_path)
    monkeypatch.setattr(rendering, "TEMPLATES_DIR", temp_dir)

    template_path = os.path.join(temp_dir, "protocol_template.docx")
    with open(template_path, "w") as f:
        f.write("Not a zip file. Invalid document structure.")

    with pytest.raises(TemplateRenderingError) as exc_info:
        rendering.load_docx_template()
    assert "Template file is invalid or corrupt" in str(exc_info.value)


def test_export_protocol_template_unavailable_integration(
    client, tmp_path, monkeypatch
):
    """
    Assert that when the controlled template is absent, a format=docx export request
    returns a structured ProblemDetails response instead of a raw 500 error.

    Requirements: PRD-SYS-001
    """
    from apps.designer import rendering

    # Isolate TEMPLATES_DIR to an empty directory
    temp_dir = str(tmp_path)
    monkeypatch.setattr(rendering, "TEMPLATES_DIR", temp_dir)

    response = client.get(
        "/api/v1/studies/study_1/export?format=docx",
        headers=get_custom_auth_headers(),
    )
    assert response.status_code == 503
    data = response.json()
    assert data["code"] == "TEMPLATE_UNAVAILABLE"
    assert "Template file is missing" in data["detail"]
    assert data["title"] == "Template Unavailable"


def test_export_protocol_as_pdf_success(client):
    """
    Verify GET /api/v1/studies/{study_id}/export with format=pdf returns a structurally
    valid PDF with correct headers and MIME types.
    """
    response = client.get(
        "/api/v1/studies/study_1/export?format=pdf",
        headers=get_custom_auth_headers(),
    )
    if response.status_code != 200:
        print("ERROR RESPONSE DETAIL:", response.text)
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"

    # Verify Content-Disposition header with safe, suggested filename
    disp = response.headers["content-disposition"]
    assert "attachment; filename=" in disp
    assert re.search(r'filename="protocol_.*_v\d+\.pdf"', disp)

    # Verify structural PDF signature (%PDF)
    content = response.content
    assert len(content) > 0
    assert content.startswith(b"%PDF")


def test_export_protocol_as_docx_success(client):
    """
    Verify GET /api/v1/studies/{study_id}/export with format=docx returns a valid
    Office Open XML document with correct headers and MIME types.
    """
    response = client.get(
        "/api/v1/studies/study_1/export?format=docx",
        headers=get_custom_auth_headers(),
    )
    if response.status_code != 200:
        print("ERROR RESPONSE DETAIL:", response.text)
    assert response.status_code == 200
    assert (
        response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Verify Content-Disposition header with safe, suggested filename
    disp = response.headers["content-disposition"]
    assert "attachment; filename=" in disp
    assert re.search(r'filename="protocol_.*_v\d+\.docx"', disp)

    # Verify OpenXML (ZIP) zipfile structure prefix (PK\x03\x04)
    content = response.content
    assert len(content) > 0
    assert content.startswith(b"PK\x03\x04")


def test_export_protocol_not_found(client):
    """
    Verify GET /api/v1/studies/{study_id}/export for non-existent study returns HTTP 404.
    """
    response = client.get(
        "/api/v1/studies/invalid_study_id_999/export?format=pdf",
        headers=get_custom_auth_headers(),
    )
    assert response.status_code == 404
    assert "Study not found" in response.json()["detail"]


def test_export_protocol_unsupported_format(client):
    """
    Verify GET /api/v1/studies/{study_id}/export with unsupported format returns HTTP 422.
    """
    response = client.get(
        "/api/v1/studies/study_1/export?format=txt",
        headers=get_custom_auth_headers(),
    )
    assert response.status_code == 422
    assert "Invalid format" in response.json()["detail"]


def test_export_protocol_invalid_output(client):
    """
    Verify GET /api/v1/studies/{study_id}/export with unsupported output returns HTTP 422.
    """
    response = client.get(
        "/api/v1/studies/study_1/export?output=invalid_section",
        headers=get_custom_auth_headers(),
    )
    assert response.status_code == 422
    assert "Invalid output" in response.json()["detail"]


def test_export_protocol_outputs_rendering(client):
    """
    Verify that all supported output sections (narrative, synopsis, soa) render successfully in PDF and DOCX.
    """
    for out in ("narrative", "synopsis", "soa", "combined"):
        for fmt in ("pdf", "docx"):
            response = client.get(
                f"/api/v1/studies/study_1/export?format={fmt}&output={out}",
                headers=get_custom_auth_headers(),
            )
            assert response.status_code == 200
            expected_mime = (
                "application/pdf"
                if fmt == "pdf"
                else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            assert response.headers["content-type"] == expected_mime


def test_export_protocol_generation_auditing(client):
    """
    Verify that successful exports are tracked in the Part 11 compliant audit trail,
    capturing caller identity, change reason, and output selection.
    """
    headers = get_custom_auth_headers(change_reason="Regulatory submission export")

    response = client.get(
        "/api/v1/studies/study_1/export?format=pdf&output=synopsis",
        headers=headers,
    )
    assert response.status_code == 200

    # Assert that one audit log event was successfully appended
    assert len(MOCK_DESIGNER_AUDIT_LOGS) == 1
    event = MOCK_DESIGNER_AUDIT_LOGS[0]
    assert event["actor"] == "123"  # matches get_custom_auth_headers user_id
    assert event["change_reason"] == "Regulatory submission export"
    assert event["study_id"] == "study_1"
    assert event["format"] == "pdf"
    assert event["output"] == "synopsis"
    assert event["type"] == "PROTOCOL_EXPORT"
    assert "timestamp" in event
    assert "id" in event


def test_export_protocol_etmf_forwarding_best_effort(client, monkeypatch):
    """
    Verify that when ETMF_STRICT_ARCHIVAL is false (default), forwarding failures do NOT block
    or invalidate a successful real-time document export.
    """
    # Force best-effort archival configuration
    monkeypatch.setenv("ETMF_FORWARDING_ENABLED", "true")
    monkeypatch.setenv("ETMF_STRICT_ARCHIVAL", "false")
    monkeypatch.setenv("ETMF_URL", "http://invalid-non-existent-etmf-url:9999")

    # The export should succeed normally even with non-existent ETMF URL (warnings logged)
    response = client.get(
        "/api/v1/studies/study_1/export?format=pdf",
        headers=get_custom_auth_headers(),
    )
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"


def test_export_protocol_etmf_forwarding_strict_failure(client, monkeypatch):
    """
    Verify that when ETMF_STRICT_ARCHIVAL is true, forwarding failures explicitly
    propagate and invalidate the export transaction.
    """
    # Force strict archival configuration
    monkeypatch.setenv("ETMF_FORWARDING_ENABLED", "true")
    monkeypatch.setenv("ETMF_STRICT_ARCHIVAL", "true")
    monkeypatch.setenv("ETMF_URL", "http://invalid-non-existent-etmf-url:9999")

    # The export should raise HTTP 500 on strict archival failure
    response = client.get(
        "/api/v1/studies/study_1/export?format=pdf",
        headers=get_custom_auth_headers(),
    )
    assert response.status_code == 500
    assert "Strict Archival Failure" in response.json()["detail"]


def test_export_protocol_unauthenticated(client):
    """
    Verify that calling the export endpoint with no gateway headers returns 401
    and "Missing gateway authentication headers" detail.
    """
    response = client.get("/api/v1/studies/study_1/export?format=pdf")
    assert response.status_code == 401
    assert "Missing gateway authentication headers" in response.json()["detail"]


def test_export_protocol_unauthorized_empty_roles(client):
    """
    Verify that calling the export endpoint with valid gateway headers but an empty/role-less
    roles value returns 403 (unauthorized/forbidden).
    """
    headers = get_custom_auth_headers(roles=" ")
    response = client.get(
        "/api/v1/studies/study_1/export?format=pdf",
        headers=headers,
    )
    assert response.status_code == 403
    # Check detail covers role missing or permission forbidden
    assert (
        "Missing role credentials" in response.json()["detail"]
        or "Forbidden" in response.json()["detail"]
    )


def _find_all_tables_in_docx(doc):
    from docx.oxml.table import CT_Tbl
    from docx.table import Table

    tables = []
    for child in doc.element.body.iter():
        if isinstance(child, CT_Tbl):
            tables.append(Table(child, doc.element.body))
    return tables


def test_render_protocol_to_docx_combined_structure():
    """
    Assert that combined DOCX output has the correct synopsis, narrative, and SoA table structure.
    """
    import io

    from docx import Document

    from apps.designer.rendering import render_protocol_to_docx

    from .test_protocol_render import get_sample_rendered_document

    doc_view = get_sample_rendered_document()
    result = render_protocol_to_docx(doc_view, "combined")

    assert result.content is not None
    assert result.filename == "protocol_study_test_v2.docx"
    assert (
        result.media_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Parse with python-docx
    doc = Document(io.BytesIO(result.content))

    # Extract all paragraph texts
    all_paras = [p.text for p in doc.paragraphs]
    full_text = "\n".join(all_paras)

    # Assert synopsis and narrative elements exist in paragraphs
    assert "CLINICAL STUDY PROTOCOL" in full_text
    assert "Sponsor Corp" in full_text
    assert "Introduction Section Title" in full_text
    assert "Intro para text" in full_text

    # Assert SoA table structure recursively
    all_tables = _find_all_tables_in_docx(doc)
    assert len(all_tables) >= 1
    soa_table = all_tables[0]

    # Assert row and column structure
    # Header 1, Header 2 + 1 activity row = 3 rows
    assert len(soa_table.rows) == 3
    # 1 activity column + 1 encounter column = 2 columns
    assert len(soa_table.columns) == 2

    # Assert cell content
    assert "Activity / Procedure" in soa_table.cell(0, 0).text
    assert "Screening" in soa_table.cell(0, 1).text
    assert "Visit 1" in soa_table.cell(1, 1).text
    assert "Vitals Collection" in soa_table.cell(2, 0).text
    assert "X" in soa_table.cell(2, 1).text
    assert "Vitals detail" in soa_table.cell(2, 1).text


def test_render_protocol_to_docx_gated_synopsis_only():
    """
    Assert that synopsis-only DOCX output omits narrative and SoA sections.
    """
    import io

    from docx import Document

    from apps.designer.rendering import render_protocol_to_docx

    from .test_protocol_render import get_sample_rendered_document

    doc_view = get_sample_rendered_document()
    result = render_protocol_to_docx(doc_view, "synopsis")

    doc = Document(io.BytesIO(result.content))
    full_text = "\n".join([p.text for p in doc.paragraphs])

    assert "PROTOCOL SYNOPSIS" in full_text
    # Section 2 narrative should not be populated with text
    assert "Introduction Section Title" not in full_text
    assert "Intro para text" not in full_text

    # The SoA table is omitted from view
    all_tables = _find_all_tables_in_docx(doc)
    assert "Schedule of Activities (SoA) omitted" in full_text or len(all_tables) == 0


def test_render_protocol_to_docx_gated_narrative_only():
    """
    Assert that narrative-only DOCX output omits synopsis and SoA sections.
    """
    import io

    from docx import Document

    from apps.designer.rendering import render_protocol_to_docx

    from .test_protocol_render import get_sample_rendered_document

    doc_view = get_sample_rendered_document()
    result = render_protocol_to_docx(doc_view, "narrative")

    doc = Document(io.BytesIO(result.content))
    full_text = "\n".join([p.text for p in doc.paragraphs])

    # Synopsis fields should be omitted from body text
    assert "PROTOCOL SYNOPSIS" not in full_text
    assert "Introduction Section Title" in full_text
    assert "Intro para text" in full_text
    all_tables = _find_all_tables_in_docx(doc)
    assert "Schedule of Activities (SoA) omitted" in full_text or len(all_tables) == 0


def test_render_protocol_to_docx_gated_soa_only():
    """
    Assert that SoA-only DOCX output omits synopsis and narrative sections.
    """
    import io

    from docx import Document

    from apps.designer.rendering import render_protocol_to_docx

    from .test_protocol_render import get_sample_rendered_document

    doc_view = get_sample_rendered_document()
    result = render_protocol_to_docx(doc_view, "soa")

    doc = Document(io.BytesIO(result.content))
    full_text = "\n".join([p.text for p in doc.paragraphs])

    assert "PROTOCOL SYNOPSIS" not in full_text
    assert "Introduction Section Title" not in full_text
    assert "Intro para text" not in full_text
    # SoA table should be present
    all_tables = _find_all_tables_in_docx(doc)
    assert len(all_tables) >= 1


def test_production_template_immutability_integration(client):
    """
    Assert that invoking the export endpoint does not modify the checked-in controlled template.
    """
    import hashlib

    from apps.designer.rendering import TEMPLATES_DIR

    real_template_path = os.path.join(TEMPLATES_DIR, "protocol_template.docx")
    assert os.path.exists(real_template_path)

    # Compute initial hash of the real template file
    with open(real_template_path, "rb") as f:
        initial_hash = hashlib.sha256(f.read()).hexdigest()

    # Call the export API
    response = client.get(
        "/api/v1/studies/study_1/export?format=docx",
        headers=get_custom_auth_headers(),
    )
    assert response.status_code == 200

    # Assert hash is unchanged
    with open(real_template_path, "rb") as f:
        after_hash = hashlib.sha256(f.read()).hexdigest()
    assert after_hash == initial_hash
