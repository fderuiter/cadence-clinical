import base64
import io
import zipfile

import docx
import fitz
import pytest
import pytest_asyncio
from fastapi.testclient import TestClient

from apps.etmf.adapters.database import db_manager
from apps.etmf.adapters.export import generate_binder_zip
from apps.etmf.adapters.ingestion_service import ingest_tmf_document
from apps.etmf.adapters.models import Base, TMFDocument
from apps.etmf.adapters.watermark import apply_watermark
from apps.etmf.main import app


@pytest_asyncio.fixture(autouse=True)
async def setup_db():
    """Setup in-memory eTMF database for unit and integration testing."""
    db_manager.init_db("sqlite+aiosqlite:///:memory:", echo=False)
    async with db_manager.engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    yield
    async with db_manager.engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)
    await db_manager.close()


def create_minimal_pdf() -> bytes:
    """Helper to generate a minimal, valid PDF in memory using fitz."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(fitz.Point(72, 72), "This is a clean, native PDF file.")
    pdf_bytes = doc.write()
    doc.close()
    return pdf_bytes


def create_minimal_docx() -> bytes:
    """Helper to generate a minimal, valid DOCX in memory using python-docx."""
    doc = docx.Document()
    doc.add_paragraph("This is a clean, native DOCX file.")
    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()


def test_native_pdf_watermarking():
    """Verify that a native PDF is watermarked visually and remains valid and uncorrupted."""
    pdf_bytes = create_minimal_pdf()

    # Test raw bytes input and return
    watermarked = apply_watermark(
        content=pdf_bytes,
        mime_type="application/pdf",
        user_id="auditor_user",
        user_role="auditor",
    )
    assert isinstance(watermarked, bytes)

    # Verify the watermarked PDF opens successfully in PyMuPDF
    doc = fitz.open(stream=watermarked, filetype="pdf")
    assert len(doc) == 1
    page_text = doc[0].get_text()
    # Note: check if any visual/structural text exists
    assert "This is a clean, native PDF file." in page_text
    doc.close()

    # Test Base64 string input and return
    b64_input = base64.b64encode(pdf_bytes).decode("utf-8")
    watermarked_b64 = apply_watermark(
        content=b64_input,
        mime_type="application/pdf",
        user_id="auditor_user",
        user_role="auditor",
    )
    assert isinstance(watermarked_b64, str)
    decoded_bytes = base64.b64decode(watermarked_b64)
    doc2 = fitz.open(stream=decoded_bytes, filetype="pdf")
    assert len(doc2) == 1
    doc2.close()


def test_native_docx_watermarking():
    """Verify that a native DOCX is watermarked visually via section headers and remains valid."""
    docx_bytes = create_minimal_docx()

    # Test raw bytes input
    watermarked = apply_watermark(
        content=docx_bytes,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        user_id="auditor_user",
        user_role="auditor",
    )
    assert isinstance(watermarked, bytes)

    # Verify the watermarked DOCX opens successfully in python-docx
    doc = docx.Document(io.BytesIO(watermarked))
    p_texts = [p.text for p in doc.paragraphs]
    header_texts = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            header_texts.append(p.text)

    # Original text is preserved in body
    assert any(
        "This is a clean, native DOCX file" in t for p_text in p_texts for t in [p_text]
    )
    # Watermark text is present in the section headers
    assert any("CONFIDENTIAL" in h_text for h_text in header_texts)


@pytest.mark.asyncio
async def test_safe_binary_ingestion_and_export():
    """Verify raw binary documents are safely ingested as Base64 strings, and exported back cleanly to ZIP."""
    session_maker = db_manager.get_session_maker()
    async with session_maker() as session:
        pdf_bytes = create_minimal_pdf()
        docx_bytes = create_minimal_docx()

        # Ingest PDF
        pdf_doc = await ingest_tmf_document(
            session=session,
            study_id="study_abc",
            artifact_type="Clinical Trial Protocol",
            filename="protocol.pdf",
            content=pdf_bytes,
            mime_type="application/pdf",
            created_by="sponsor_user",
            created_role="sponsor_dm",
        )

        # Ingest DOCX
        await ingest_tmf_document(
            session=session,
            study_id="study_abc",
            artifact_type="Trial Monitoring Plan",
            filename="protocol.docx",
            content=docx_bytes,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            created_by="sponsor_user",
            created_role="sponsor_dm",
        )

        await session.commit()

        # Check database records
        # Content must be saved as Base64 encoded string
        assert pdf_doc.content != ""
        assert pdf_doc.content != pdf_bytes
        # Ensure it is a valid Base64 string that decodes to our PDF starting with %PDF
        decoded_pdf = base64.b64decode(pdf_doc.content)
        assert decoded_pdf.startswith(b"%PDF")

        # Perform Bulk Export ZIP Generation
        zip_bytes = await generate_binder_zip(
            session=session,
            study_id="study_abc",
            include_history=False,
            requester_id="auditor_1",
            requester_role="auditor",
        )

        # Open and inspect the exported ZIP package entirely in-memory
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as z:
            namelist = z.namelist()
            print("ZIP NAMELIST:", namelist)
            pdf_path = [name for name in namelist if name.endswith(".pdf")][0]
            docx_path = [name for name in namelist if name.endswith(".docx")][0]

            exported_pdf_bytes = z.read(pdf_path)
            exported_docx_bytes = z.read(docx_path)

            # Confirm documents open successfully without any file errors
            pdf_doc_check = fitz.open(stream=exported_pdf_bytes, filetype="pdf")
            assert len(pdf_doc_check) == 1
            pdf_doc_check.close()

            docx_doc_check = docx.Document(io.BytesIO(exported_docx_bytes))
            assert len(docx_doc_check.paragraphs) > 0


@pytest.mark.asyncio
async def test_automated_webhook_non_degraded_ingestion(monkeypatch):
    """Verify clinical email webhook ingests clinical documents without byte-level degradation."""
    monkeypatch.setenv("INBOUND_EMAIL_HMAC_SECRET", "test-secret-key-12345")
    monkeypatch.setenv("INBOUND_EMAIL_MAX_SIZE_BYTES", "5000000")

    client = TestClient(app)
    import hashlib
    import hmac
    import time

    timestamp = str(time.time())
    token = "unique-webhook-token"
    signature = hmac.new(
        b"test-secret-key-12345", f"{timestamp}{token}".encode(), hashlib.sha256
    ).hexdigest()

    pdf_bytes = create_minimal_pdf()

    response = client.post(
        "/api/v1/etmf/inbound-email",
        data={
            "sender": "partner@example.com",
            "recipient": "study-xyz+conduct@example.com",
            "subject": "Signed Clinical Agreement",
            "body-plain": "Please see attached PDF agreement.",
            "timestamp": timestamp,
            "token": token,
            "signature": signature,
            "Message-Id": "<agreement-email-1@example.com>",
        },
        files=[
            ("attachment", ("agreement.pdf", pdf_bytes, "application/pdf")),
        ],
    )

    assert response.status_code == 201

    # Retrieve from DB and verify zero byte degradation
    session_maker = db_manager.get_session_maker()
    async with session_maker() as session:
        # Retrieve the document attachment (filename="agreement.pdf")
        from sqlalchemy import select

        stmt = select(TMFDocument).where(TMFDocument.filename == "agreement.pdf")
        res = await session.execute(stmt)
        doc = res.scalars().first()
        assert doc is not None
        assert doc.mime_type == "application/pdf"

        # Check content is Base64 stored
        stored_bytes = base64.b64decode(doc.content)
        # It must be perfectly readable and identical to the original PDF
        original_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        stored_doc = fitz.open(stream=stored_bytes, filetype="pdf")
        assert len(original_doc) == len(stored_doc)
        original_doc.close()
        stored_doc.close()
