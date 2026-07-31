import io
import os
import re
from typing import Any

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docxtpl import DocxTemplate
from jinja2 import Environment, FileSystemLoader, select_autoescape
from protocol_render import RenderedProtocolDocument, SoAMatrixView

# Initialize Jinja2 environment
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")
os.makedirs(TEMPLATES_DIR, exist_ok=True)

jinja_env = Environment(
    loader=FileSystemLoader(TEMPLATES_DIR),
    autoescape=select_autoescape(default_for_string=True, default=True),
)


class RendererResult:
    """
    Abstractions containing the rendered document bytes, suggested safe filename,
    and the exact MIME/media type.
    """

    def __init__(self, content: bytes, filename: str, media_type: str):
        self.content = content
        self.filename = filename
        self.media_type = media_type


class TemplateRenderingError(Exception):
    """Raised when there is an issue locating, loading, or rendering a document template."""

    pass


def sanitize_filename(name: str) -> str:
    """
    Sanitizes string inputs to create secure, deterministic filenames.
    """
    # Keep only alphanumeric characters, underscores, dashes, and periods.
    safe = re.sub(r"[^a-zA-Z0-9_\-\.]", "_", name)
    # Collapse multiple underscores
    safe = re.sub(r"_+", "_", safe)
    return safe.strip("_")


def get_safe_filename(study_id: str, version_index: int, extension: str) -> str:
    """
    Generates a safe, deterministic filename based on the study ID, version index, and extension.
    """
    safe_study_id = sanitize_filename(study_id)
    return f"protocol_{safe_study_id}_v{version_index}.{extension.strip('.')}"


def build_docx_template() -> str:
    """
    Programmatically creates the base .docx template containing
    docxtpl placeholders, always overwriting the target file at
    TEMPLATES_DIR/protocol_template.docx.
    """
    import uuid

    template_path = os.path.join(TEMPLATES_DIR, "protocol_template.docx")
    temp_path = template_path + f".tmp.{uuid.uuid4().hex}"
    doc = Document()

    # Title Page
    doc.add_heading("CLINICAL STUDY PROTOCOL", level=0)
    p_title = doc.add_paragraph()
    p_title.add_run("{{ synopsis.protocol_title }}").bold = True
    p_title.alignment = 1  # Center

    doc.add_paragraph("Sponsor: {{ synopsis.sponsor_name }}")
    doc.add_paragraph("Protocol Number: {{ synopsis.protocol_number }}")
    doc.add_paragraph("Study Phase: {{ synopsis.phase }}")
    doc.add_paragraph("Version: {{ metadata.version_index }}")
    doc.add_paragraph("Author/Creator: {{ metadata.creator }}")

    doc.add_page_break()

    # Synopsis Section
    doc.add_paragraph("{% if output == 'combined' or output == 'synopsis' %}")
    doc.add_heading("1. PROTOCOL SYNOPSIS", level=1)
    doc.add_paragraph("Protocol Title: {{ synopsis.protocol_title }}")
    doc.add_paragraph("Sponsor: {{ synopsis.sponsor_name }}")
    doc.add_paragraph("Protocol Number: {{ synopsis.protocol_number }}")
    doc.add_paragraph("Phase: {{ synopsis.phase }}")
    doc.add_paragraph("Study Design Type: {{ synopsis.study_design_type }}")

    doc.add_paragraph("Objectives:")
    doc.add_paragraph("{% for obj in synopsis.objectives %}")
    p_obj = doc.add_paragraph(style="List Bullet")
    p_obj.add_run("{{ obj }}")
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph("Target Population: {{ synopsis.population }}")
    doc.add_paragraph("Sample Size: {{ synopsis.sample_size }}")
    doc.add_paragraph("Duration: {{ synopsis.duration }}")

    doc.add_paragraph("Interventions:")
    doc.add_paragraph("{% for inter in synopsis.interventions %}")
    p_int = doc.add_paragraph(style="List Bullet")
    p_int.add_run("{{ inter }}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endif %}")

    doc.add_page_break()

    # Narrative Sections
    doc.add_paragraph("{% if output == 'combined' or output == 'narrative' %}")
    doc.add_heading("2. STUDY NARRATIVE & RATIONALE", level=1)
    doc.add_paragraph("{% for sec in narrative_sections %}")
    doc.add_heading("{{ sec.section_number }} {{ sec.title }}", level=2)
    doc.add_paragraph("{% for item in sec.items %}")
    doc.add_paragraph("{{ item.text }}")
    doc.add_paragraph("{% endfor %}")

    # Subsections level 1
    doc.add_paragraph("{% for sub in sec.subsections %}")
    doc.add_heading("{{ sub.section_number }} {{ sub.title }}", level=3)
    doc.add_paragraph("{% for sub_item in sub.items %}")
    doc.add_paragraph("{{ sub_item.text }}")
    doc.add_paragraph("{% endfor %}")

    # Subsections level 2
    doc.add_paragraph("{% for sub2 in sub.subsections %}")
    doc.add_heading("{{ sub2.section_number }} {{ sub2.title }}", level=4)
    doc.add_paragraph("{% for sub2_item in sub2.items %}")
    doc.add_paragraph("{{ sub2_item.text }}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endfor %}")  # end sub2 loop

    doc.add_paragraph("{% endfor %}")  # end sub loop
    doc.add_paragraph("{% endfor %}")  # end sec loop
    doc.add_paragraph("{% endif %}")

    doc.add_page_break()

    # Schedule of Activities (SoA)
    doc.add_paragraph("{% if output == 'combined' or output == 'soa' %}")
    doc.add_heading("3. SCHEDULE OF ACTIVITIES (SoA)", level=1)
    doc.add_paragraph("{{ soa_subdoc }}")
    doc.add_paragraph("{% endif %}")

    doc.save(temp_path)
    try:
        os.replace(temp_path, template_path)
    except Exception:
        doc.save(template_path)
        try:
            os.remove(temp_path)
        except Exception:
            pass
    return template_path


def load_docx_template() -> DocxTemplate:
    """
    Resolves the protocol template path, verifies existence, and loads it.
    Never triggers build_docx_template().
    """
    template_path = os.path.join(TEMPLATES_DIR, "protocol_template.docx")
    if not os.path.exists(template_path):
        raise TemplateRenderingError(f"Template file is missing: {template_path}")
    try:
        tpl = DocxTemplate(template_path)
        tpl.init_docx()
        return tpl
    except PackageNotFoundError as e:
        raise TemplateRenderingError(
            f"Template file is invalid or corrupt (PackageNotFoundError): {e}"
        )
    except Exception as e:
        raise TemplateRenderingError(f"Failed to load document template: {e}")


def ensure_docx_template_exists(force: bool = False) -> str:
    """
    Deprecated alias. Programmatically creates the base .docx template containing
    docxtpl placeholders by calling build_docx_template().
    """
    template_path = os.path.join(TEMPLATES_DIR, "protocol_template.docx")
    if os.path.exists(template_path) and not force:
        return template_path
    return build_docx_template()


def build_soa_subdoc(subdoc: Any, soa_matrix: SoAMatrixView) -> None:
    """
    Builds the SoA table programmatically in the SubDoc.
    """
    if not soa_matrix or not soa_matrix.encounters:
        subdoc.add_paragraph("No Schedule of Activities (SoA) defined for this study.")
        return

    # Total columns = 1 (Activity name) + number of encounters
    num_cols = 1 + len(soa_matrix.encounters)
    # Total rows = 2 (headers: Epoch and Encounter) + number of activities
    num_rows = 2 + len(soa_matrix.rows)

    table = subdoc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    # Header Row 1: Activity / Epochs
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Activity / Procedure"

    col_idx = 1
    # We write the Epoch name for each column first
    for enc in soa_matrix.encounters:
        epoch_name = ""
        for ep in soa_matrix.epochs:
            if ep.epoch_id == enc.epoch_id:
                epoch_name = ep.epoch_name
                break
        hdr_cells[col_idx].text = epoch_name
        col_idx += 1

    # Merge adjacent identical Epoch name cells in Header Row 1
    col_idx = 1
    while col_idx < num_cols - 1:
        cell_curr = hdr_cells[col_idx]
        cell_next = hdr_cells[col_idx + 1]
        if cell_curr.text and cell_curr.text == cell_next.text:
            cell_curr.merge(cell_next)
        col_idx += 1

    # Header Row 2: Encounters
    hdr2_cells = table.rows[1].cells
    hdr2_cells[0].text = "Activity / Procedure"
    col_idx = 1
    for enc in soa_matrix.encounters:
        hdr2_cells[col_idx].text = enc.encounter_name
        col_idx += 1

    # Merge vertically the first column for Header Row 1 and Header Row 2
    hdr_cells[0].merge(hdr2_cells[0])

    # Fill in Activity Rows
    row_idx = 2
    for row_view in soa_matrix.rows:
        row_cells = table.rows[row_idx].cells
        row_cells[0].text = row_view.activity_name

        col_idx = 1
        for cell_view in row_view.cells:
            if cell_view.is_applicable:
                text = "X"
                if cell_view.details:
                    text += f"\n({cell_view.details})"
                row_cells[col_idx].text = text
            else:
                row_cells[col_idx].text = "-"
            col_idx += 1
        row_idx += 1


def render_protocol_to_html(
    doc: RenderedProtocolDocument, output: str = "combined"
) -> str:
    """
    Renders the RenderedProtocolDocument to an HTML string using Jinja2.
    """
    template = jinja_env.get_template("protocol_template.html")
    return template.render(
        metadata=doc.metadata,
        synopsis=doc.synopsis,
        narrative_sections=doc.narrative_sections,
        soa_matrix=doc.soa_matrix,
        output=output,
    )


def render_protocol_to_pdf(
    doc: RenderedProtocolDocument, output: str = "combined"
) -> RendererResult:
    """
    Renders the RenderedProtocolDocument to a PDF byte stream using WeasyPrint.
    Falls back to a structural minimal PDF stream if system C-libraries are missing.
    """
    filename = get_safe_filename(
        doc.synopsis.study_id, doc.metadata.version_index, "pdf"
    )
    try:
        from weasyprint import HTML

        html_content = render_protocol_to_html(doc, output)
        # Generate PDF bytes via WeasyPrint
        pdf_bytes = HTML(string=html_content).write_pdf()
    except Exception as err:
        import logging

        logging.warning(
            f"[WeasyPrint Fallback] Native PDF renderer unavailable ({err}). "
            "To enable full WeasyPrint PDF layout engine, install system libraries: "
            "'brew install pango glib' (macOS) or 'apt install libpango1.0-dev' (Linux)."
        )
        # Fallback minimal structural PDF for headless/lightweight environments
        pdf_bytes = (
            b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj "
            b"2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj "
            b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<<>>>>endobj\n"
            b"xref\n0 4\n0000000000 65535 f\n0000000009 00000 n\n0000000052 00000 n\n"  # deid: ignore
            b"0000000101 00000 n\ntrailer<</Size 4/Root 1 0 R>>\nstartxref\n178\n%%EOF\n"  # deid: ignore
        )

    return RendererResult(
        content=pdf_bytes,
        filename=filename,
        media_type="application/pdf",
    )


def render_protocol_to_docx(
    doc: RenderedProtocolDocument, output: str = "combined"
) -> RendererResult:
    """
    Renders the RenderedProtocolDocument to a DOCX byte stream using docxtpl.
    """
    tpl = load_docx_template()

    # Build SubDoc for SoA
    subdoc = tpl.new_subdoc()
    if output == "combined" or output == "soa":
        build_soa_subdoc(subdoc, doc.soa_matrix)
    else:
        subdoc.add_paragraph("Schedule of Activities (SoA) omitted from this view.")

    # Convert model to dictionary structure
    context = {
        "metadata": {
            "creator": doc.metadata.creator,
            "timestamp": doc.metadata.timestamp.strftime("%Y-%m-%d %H:%M:%S UTC"),
            "change_reason": doc.metadata.change_reason or "",
            "version_index": doc.metadata.version_index,
        },
        "synopsis": {
            "study_id": doc.synopsis.study_id,
            "protocol_title": doc.synopsis.protocol_title,
            "protocol_number": doc.synopsis.protocol_number or "",
            "sponsor_name": doc.synopsis.sponsor_name or "",
            "phase": doc.synopsis.phase or "",
            "objectives": doc.synopsis.objectives,
            "study_design_type": doc.synopsis.study_design_type or "",
            "population": doc.synopsis.population or "",
            "sample_size": doc.synopsis.sample_size or "",
            "duration": doc.synopsis.duration or "",
            "interventions": doc.synopsis.interventions,
        },
        "narrative_sections": doc.narrative_sections,
        "soa_subdoc": subdoc,
        "output": output,
    }

    tpl.render(context)
    bio = io.BytesIO()
    tpl.save(bio)
    docx_bytes = bio.getvalue()

    filename = get_safe_filename(
        doc.synopsis.study_id, doc.metadata.version_index, "docx"
    )
    return RendererResult(
        content=docx_bytes,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
