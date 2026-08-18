"""Watermarking helper module for eTMF Service providing secure document watermarking."""

import base64
import json
from datetime import UTC, datetime


def apply_watermark(
    content: str | bytes, mime_type: str, user_id: str, user_role: str
) -> str | bytes:
    """Applies a secure, attributable watermark to the given document content.

    Args:
        content (str | bytes): The document content to watermark.
        mime_type (str): The MIME type of the document.
        user_id (str): The ID of the requester.
        user_role (str): The role/roles of the requester.

    Returns:
        str | bytes: The watermarked content.
    """
    now_utc = datetime.now(UTC).isoformat()
    marker = "CONFIDENTIAL — Auditor Copy"
    watermark_msg = (
        f"{marker} | Access by: {user_id} ({user_role}) | UTC Time: {now_utc}"
    )

    mime_lower = mime_type.lower().strip()
    is_pdf = "pdf" in mime_lower
    is_docx = "wordprocessingml" in mime_lower or "docx" in mime_lower

    input_is_bytes = isinstance(content, bytes)

    # 1. Native Binary Watermarking (PDF / DOCX)
    if is_pdf or is_docx:
        raw_bytes: bytes | None = None
        is_b64 = False

        if isinstance(content, bytes):
            raw_bytes = content
        else:
            try:
                decoded = base64.b64decode(content)
                if (is_pdf and decoded.startswith(b"%PDF")) or (
                    is_docx and decoded.startswith(b"PK\x03\x04")
                ):
                    raw_bytes = decoded
                    is_b64 = True
            except Exception:
                pass

        if raw_bytes is not None and isinstance(raw_bytes, bytes):
            if is_pdf:
                try:
                    import fitz

                    doc = fitz.open(stream=raw_bytes, filetype="pdf")
                    for page in doc:
                        rect = page.rect
                        point = fitz.Point(36, rect.height - 36)
                        page.insert_text(
                            point, watermark_msg, fontsize=9, color=(0.7, 0.1, 0.1)
                        )
                    watermarked_bytes = doc.write()
                    doc.close()

                    if is_b64:
                        return base64.b64encode(watermarked_bytes).decode("utf-8")
                    if input_is_bytes:
                        return watermarked_bytes
                    return watermarked_bytes.decode("utf-8", errors="ignore")
                except Exception:
                    pass

            elif is_docx:
                try:
                    import io

                    import docx

                    doc = docx.Document(io.BytesIO(raw_bytes))
                    for section in doc.sections:
                        header = section.header
                        p = (
                            header.paragraphs[0]
                            if header.paragraphs
                            else header.add_paragraph()
                        )
                        p.text = f"{watermark_msg}\n{p.text}".strip()

                    out_io = io.BytesIO()
                    doc.save(out_io)
                    watermarked_bytes = out_io.getvalue()

                    if is_b64:
                        return base64.b64encode(watermarked_bytes).decode("utf-8")
                    if input_is_bytes:
                        return watermarked_bytes
                    return watermarked_bytes.decode("utf-8", errors="ignore")
                except Exception:
                    pass

    # 2. Text formats & Fallback
    if isinstance(content, bytes):
        content_str = content.decode("utf-8", errors="ignore")
    else:
        content_str = str(content)

    if "json" in mime_lower:
        try:
            parsed = json.loads(content_str)
            if isinstance(parsed, dict):
                parsed["_watermark"] = {
                    "marker": marker,
                    "accessed_by": user_id,
                    "role": user_role,
                    "timestamp": now_utc,
                }
                result_str = json.dumps(parsed, indent=2)
            else:
                result_str = content_str
        except Exception:
            result_str = content_str

    elif "xml" in mime_lower or "html" in mime_lower:
        comment = f"\n<!-- {watermark_msg} -->"
        result_str = content_str + comment

    elif "csv" in mime_lower:
        row = f"\n# {watermark_msg}"
        result_str = content_str + row

    else:
        fallback_block = (
            f"\n\n--- WATERMARK ---\n"
            f"CONFIDENTIAL — Auditor Copy\n"
            f"Accessed by: {user_id} ({user_role})\n"
            f"UTC Timestamp: {now_utc}\n"
            f"------------------\n"
        )
        result_str = content_str + fallback_block

    if input_is_bytes:
        return result_str.encode("utf-8")
    return result_str


__all__ = ["apply_watermark"]
